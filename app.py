import streamlit as st
import psycopg2
import mysql.connector
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from db_connection import get_doc_db_connection

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# def add_table_of_contents(doc):
#     """
#     یک فیلد TOC استاندارد اضافه می‌کند.
#     Word بعد از باز شدن سند با F9 یا Update Field شماره صفحات را می‌سازد.
#     """
#     p = doc.add_paragraph()                 # پاراگراف خالی برای فیلد
#     run = p.add_run()

#     # آغاز فیلد
#     fld_begin = OxmlElement('w:fldChar')
#     fld_begin.set(qn('w:fldCharType'), 'begin')
#     run._r.append(fld_begin)

#     # دستور TOC – تیترهای Heading1 تا Heading3، لینک‌دار (\\h) و بدون خطوط فرمت (\\z)
#     instr = OxmlElement('w:instrText')
#     instr.text = r'TOC \o "1-3" \h \z \u'   # \u جدول را راست‌چین می‌کند
#     run._r.append(instr)

#     # جداکننده
#     fld_sep = OxmlElement('w:fldChar')
#     fld_sep.set(qn('w:fldCharType'), 'separate')
#     run._r.append(fld_sep)

#     # متن موقتی که قبل از به‌روزرسانی دیده می‌شود
#     dummy = OxmlElement('w:t')
#     dummy.text = "برای به‌روزرسانی فهرست مطالب، کلید F9 را بزنید."
#     run._r.append(dummy)

#     # پایان فیلد
#     fld_end = OxmlElement('w:fldChar')
#     fld_end.set(qn('w:fldCharType'), 'end')
#     run._r.append(fld_end)

def add_table_of_contents(doc):
    """
    اضافه کردن فهرست مطالب با جهت راست به چپ و فونت سفارشی.
    """

    # پاراگراف فهرست مطالب
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # راست‌چین

    # جهت پاراگراف (راست به چپ)
    p_para = p._p
    pPr = p_para.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

    run = p.add_run()

    # فونت فارسی و انگلیسی
    font = run.font
    font.name = 'Times New Roman'     # این فقط برای متن لاتین مؤثر است
    font.size = Pt(12)
    rPr = run._element.get_or_add_rPr()

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')         # فونت انگلیسی
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')         # برای متن‌های لاتین
    rFonts.set(qn('w:cs'), 'Complex')                    # فونت فارسی (Complex)
    rPr.append(rFonts)

    # آغاز فیلد
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    # دستور TOC
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')  # حفظ فاصله‌ها
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    # جداکننده
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    # متن موقتی
    dummy = OxmlElement('w:t')
    dummy.text = "برای به‌روزرسانی فهرست مطالب، کلید F9 را بزنید."
    run._r.append(dummy)

    # پایان فیلد
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

# بارگذاری CSS برای راست به چپ (RTL)
st.markdown(
    """
    <style>
    @font-face {
        font-family: 'Vazir';
        src: url('/app/static/media/Vazirmatn-Regular.woff2') format('woff2');
        font-weight: normal;
        font-style: normal;
    }
    body {
        direction: rtl;
        text-align: right;        
    }

    .rtl {
        direction: rtl;
        text-align: right;
        font-family: 'Vazir';
        
    }
    input::placeholder {
        text-align: right;
        direction: rtl;
        margin-top: -10px;
        font-family: 'Vazir'; 
    }
    .custom-label {
        direction: rtl;
        text-align: right;
        margin-bottom: -10px; 
        font-family: 'Vazir'; 
    }
    .st-emotion-cache-162xg8y {
        margin-top: -10px;  
        margin-bottom: -10px; 
        padding: 0px; 
    }
    .info{
        font-family:Vazir;
        background-color:#e1e1e1;
        padding:10px;
        border-radius: 10px;
        margin-bottom: 10px;
    }
    .success{
        font-family:Vazir;
        background-color:#eebbff;
        padding:10px;
        border-radius: 10px;
        margin-bottom: 10px;
    }
    </style>
    """, unsafe_allow_html=True
)

def init_session_state(selected_table):
    if selected_table not in st.session_state:
        st.session_state[selected_table] = {}

def render_table_form(selected_table):
    st.session_state.selected_table = selected_table
    conn = st.session_state.connection
    cursor = conn.cursor()

    if st.session_state.selected_table:
        st.markdown("### اطلاعات جدول")
        st.text_input("نام جدول:", value=st.session_state.selected_table, disabled=True)
        table_fa_name = st.text_input("نام فارسی جدول:")
        table_description = st.text_area("توضیح جدول:")
        st.markdown("---")
        st.markdown("### اطلاعات فیلدها")
        columns = get_table_columns(cursor, st.session_state.selected_table)
        for col in columns:
            col_name = col[0]
            data_type = col[1]
            st.markdown(f"**{col_name}: {data_type}**")
            st.text_input(f"نام فارسی فیلد ({col_name}):", key=f"{selected_table}_{col_name}_fa")
            st.text_area(f"توضیح فیلد ({col_name}):", key=f"{selected_table}_{col_name}_desc")
            st.markdown("---")


def connect_to_db(db_type, ip, port, user, password, db_name):
    try:
        if db_type == "PostgreSQL":
            conn = psycopg2.connect(
                host=ip, port=port, user=user, password=password, dbname=db_name
            )
            cursor = conn.cursor()
            cursor.execute("""SELECT table_name FROM information_schema.tables 
                              WHERE table_schema='public'""")
        elif db_type == "MySQL":
            conn = mysql.connector.connect(
                host=ip, port=port, user=user, password=password, database=db_name
            )
            cursor = conn.cursor()
            cursor.execute("SHOW TABLES")
            
        else:
            raise Exception("Unknown database type")

        tables = cursor.fetchall()
        cursor.close()

        # ذخیره اتصال در session_state
        st.session_state.connection = conn
        st.session_state.db_name = db_name
        st.session_state.db_type = db_type
        return tables

    except Exception as e:
        st.error(f"اتصال به دیتابیس با خطا مواجه شد:\n{e}")
        return []

def ensure_main_connection():
    try:
        if "connection" in st.session_state:
            # تست اتصال با اجرای یک کوئری ساده
            st.session_state.connection.ping(reconnect=True, attempts=1, delay=0) if st.session_state.db_type == "MySQL" \
                else st.session_state.connection.cursor().execute("SELECT 1")
        else:
            # اطلاعات اتصال رو از session بردار و دوباره وصل شو
            required_keys = ["db_type", "ip", "port", "user", "password", "db_name"]
            if all(k in st.session_state for k in required_keys):
                connect_to_db(
                    st.session_state.db_type,
                    st.session_state.ip,
                    st.session_state.port,
                    st.session_state.user,
                    st.session_state.password,
                    st.session_state.db_name
                )
            else:
                st.error("اطلاعات اتصال ناقص است.")
                return False
        return True
    except Exception as e:
        st.error(f"خطا در بررسی اتصال: {e}")
        return False


def db_connection_form():
    
    st.markdown('<h2 style="font-family:Vazir;" class="rtl">🧾 فرم اتصال به دیتابیس</h2>', unsafe_allow_html=True)
    st.markdown('<h5 style="font-family:Vazir;" class="rtl">لطفاً اطلاعات دیتابیس خود را وارد کنید:</h5>', unsafe_allow_html=True)
    # st.write("لطفاً اطلاعات دیتابیس خود را وارد کنید:")

    st.markdown('<p class="custom-label">نوع دیتابیس</p>', unsafe_allow_html=True)
    db_type = st.selectbox("", ["MySQL", "PostgreSQL"], label_visibility="hidden")
    st.markdown('<p class="custom-label">IP دیتابیس</p>', unsafe_allow_html=True)
    ip = st.text_input("", "127.0.0.1", label_visibility="hidden")
    st.markdown('<p class="custom-label">پورت دیتابیس</p>', unsafe_allow_html=True)
    port = st.text_input("", "3307" if db_type == "MySQL" else "5432", label_visibility="hidden")
    st.markdown('<p class="custom-label">نام کاربری</p>', unsafe_allow_html=True)
    user = st.text_input("یوزر", label_visibility="hidden")
    st.markdown('<p class="custom-label">رمز عبور</p>', unsafe_allow_html=True)
    password = st.text_input("", type="password", label_visibility="hidden")
    st.markdown('<p class="custom-label">نام دیتابیس</p>', unsafe_allow_html=True)
    db_name = st.text_input("نام دیتابیس", label_visibility="hidden")

    if st.button("اتصال به دیتابیس"):
        if not db_name:
            #st.error("لطفاً نام دیتابیس را وارد کنید.")
            st.markdown('<p class="custom-label info">📛 لطفاً نام دیتابیس را وارد کنید.</p>', unsafe_allow_html=True)
        else:
            tables = connect_to_db(db_type, ip, port, user, password, db_name)
            excluded_tables = [
                'file_extensions', 'file_folders', 'file_servers', 'files_copy1',
                'pg_stat_statements','pg_stat_statements_info'
            ]

            if tables:
                st.session_state.db_connected = True
                filtered_tables = sorted(
                    [table for table in tables if table[0] not in excluded_tables],
                    key = lambda x: x[0]
                )

                st.session_state.tables = filtered_tables
                print(st.session_state.tables)
                st.success("اتصال به دیتابیس برقرار شد!")
                db_tables_page()

def set_font(paragraph, font_name="B Nazanin", font_size=14):
    run = paragraph.runs[0]
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

def is_persian_or_digit(text):
    for ch in text:
        if ch.isdigit():
            continue
        if '\u0600' <= ch <= '\u06FF' or '\u0750' <= ch <= '\u077F':
            continue
        return False
    return True


def set_bidi_alignment(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_run_language(run, lang="fa-IR"):
    r = run._element
    rPr = r.get_or_add_rPr()
    lang_element = OxmlElement('w:lang')
    lang_element.set(qn('w:val'), lang)
    lang_element.set(qn('w:eastAsia'), lang)
    lang_element.set(qn('w:bidi'), lang)
    rPr.append(lang_element)

def set_cell_font(cell, text, font_name="B Nazanin", font_size=12, align_center=False, bold=False, rtl=False):
    cell.text = ""
    para = cell.paragraphs[0]
    
    # وسط‌چین
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif rtl:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_bidi_alignment(para)

    run = para.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def get_table_columns(cursor, table_name):
    if isinstance(table_name, (tuple, list)):
        table_name = table_name[0]

    db_type = st.session_state.get("db_type", "MySQL")

    if db_type == "PostgreSQL":
        # گرفتن اطلاعات ستون‌ها
        cursor.execute("""
            SELECT 
                column_name,
                data_type,
                is_nullable,
                column_default
            FROM information_schema.columns
            WHERE table_name = %s AND table_schema = 'public'
            ORDER BY ordinal_position
        """, (table_name,))
        columns = cursor.fetchall()

        # گرفتن ارتباطات Foreign Key
        cursor.execute("""
            SELECT 
                kcu.column_name,
                ccu.table_name AS referenced_table
            FROM 
                information_schema.table_constraints AS tc
                JOIN information_schema.key_column_usage AS kcu
                  ON tc.constraint_name = kcu.constraint_name
                  AND tc.table_schema = kcu.table_schema
                JOIN information_schema.constraint_column_usage AS ccu
                  ON ccu.constraint_name = tc.constraint_name
                  AND ccu.constraint_schema = tc.constraint_schema
            WHERE tc.constraint_type = 'FOREIGN KEY' AND tc.table_name = %s;
        """, (table_name,))
        fk_map = dict(cursor.fetchall())

        # ترکیب اطلاعات ستون‌ها با FK
        result = []
        for col in columns:
            col_name = col[0]
            data_type = col[1]
            is_nullable = col[2]
            default = col[3]
            referenced_table = fk_map.get(col_name)
            result.append((col_name, data_type, is_nullable, default, referenced_table))
        return result

    else:  # MySQL
        cursor.execute("""
            SELECT 
                c.column_name,
                c.data_type,
                c.is_nullable,
                c.column_default,
                kcu.REFERENCED_TABLE_NAME
            FROM information_schema.columns c
            LEFT JOIN information_schema.KEY_COLUMN_USAGE kcu
                ON c.table_name = kcu.table_name
                AND c.column_name = kcu.column_name
                AND c.table_schema = kcu.table_schema
            WHERE c.table_name = %s
                AND c.table_schema = %s
            ORDER BY c.ordinal_position
        """, (table_name, st.session_state.db_name))
        return cursor.fetchall()

    if isinstance(table_name, (tuple, list)):
        table_name = table_name[0]

    db_type = st.session_state.get("db_type", "MySQL")

    if db_type == "PostgreSQL":
        cursor.execute("""
            SELECT 
                column_name,
                data_type,
                is_nullable,
                column_default
            FROM information_schema.columns
            WHERE table_name = %s AND table_schema = 'public'
            ORDER BY ordinal_position
        """, (table_name,))
        columns = cursor.fetchall()
        # برای PostgreSQL ستون foreign key رو جدا بگیریم
        cursor.execute("""
            SELECT 
                kcu.column_name,
                ccu.table_name AS referenced_table
            FROM 
                information_schema.table_constraints AS tc
                JOIN information_schema.key_column_usage AS kcu
                  ON tc.constraint_name = kcu.constraint_name
                  AND tc.table_schema = kcu.table_schema
                JOIN information_schema.constraint_column_usage AS ccu
                  ON ccu.constraint_name = tc.constraint_name
                  AND ccu.table_schema = tc.table_schema
            WHERE tc.constraint_type = 'FOREIGN KEY' AND tc.table_name = %s;
        """, (table_name,))
        fk_map = dict(cursor.fetchall())

        result = []
        for col in columns:
            col_name = col[0]
            data_type = col[1]
            is_nullable = col[2]
            default = col[3]
            referenced_table = fk_map.get(col_name)
            result.append((col_name, data_type, is_nullable, default, referenced_table))
        return result

    else:  # MySQL
        cursor.execute("""
            SELECT 
                c.column_name,
                c.data_type,
                c.is_nullable,
                c.column_default,
                kcu.REFERENCED_TABLE_NAME
            FROM information_schema.columns c
            LEFT JOIN information_schema.KEY_COLUMN_USAGE kcu
                ON c.table_name = kcu.table_name
                AND c.column_name = kcu.column_name
                AND c.table_schema = kcu.table_schema
            WHERE c.table_name = %s
                AND c.table_schema = %s
            ORDER BY c.ordinal_position
        """, (table_name, st.session_state.db_name))
        return cursor.fetchall()


def get_foreign_key_join_clauses(cursor, table_name):
    if isinstance(table_name, (tuple, list)):
        table_name = table_name[0]

    db_type = st.session_state.get("db_type")
    db_name = st.session_state.db_name

    if db_type == "PostgreSQL":
        cursor.execute("""
            SELECT
                tc.table_name AS foreign_table,
                ccu.table_name AS primary_table,
                tc.table_name || '.' || kcu.column_name || ' = ' || 
                ccu.table_name || '.' || ccu.column_name AS join_clause
            FROM 
                information_schema.table_constraints AS tc 
                JOIN information_schema.key_column_usage AS kcu
                  ON tc.constraint_name = kcu.constraint_name
                 AND tc.table_schema = kcu.table_schema
                JOIN information_schema.constraint_column_usage AS ccu
                  ON ccu.constraint_name = tc.constraint_name
                 AND ccu.table_schema = tc.table_schema
            WHERE tc.constraint_type = 'FOREIGN KEY'
              AND (tc.table_name = %s OR ccu.table_name = %s)
              AND tc.table_schema = 'public'
        """, (table_name, table_name))

    elif db_type == "MySQL":
        cursor.execute("""
            SELECT 
                table_name AS foreign_table,
                referenced_table_name AS primary_table,
                CONCAT(table_name, '.', column_name, ' = ', referenced_table_name, '.', referenced_column_name) AS join_clause
            FROM information_schema.KEY_COLUMN_USAGE
            WHERE table_schema = %s
              AND (table_name = %s OR referenced_table_name = %s)
              AND referenced_table_name IS NOT NULL
        """, (db_name, table_name, table_name))

    return cursor.fetchall()


def db_tables_page():
    db_type = st.session_state.get("db_type")
    db_name = st.session_state.db_name
    #st.title("نوع دیتابیس: " + str(db_type) + " | نام دیتابیس: " + str(db_name))
    st.title(f"نوع دیتابیس: {db_type}")
    st.title(f"نام دیتابیس: {db_name}")

    if st.session_state.get("db_connected"):

        if st.button("ساخت مستند دیتابیس به فارسی"):
            st.write("در حال ساخت مستند دیتابیس به زبان فارسی...")

            conn = st.session_state.connection
            cursor = conn.cursor()
            db_name = st.session_state.get("db_name", "database")

            document = Document()

            # جلد
            cover_para = document.add_paragraph()
            cover_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_bidi_alignment(cover_para)
            cover_run = cover_para.add_run('مستند بانک اطلاعاتی\n\n\n\n\n\n\nنام فارسی دیتابیس\nنام سازمان')
            cover_run.font.name = 'B Titr'
            cover_run._element.rPr.rFonts.set(qn('w:cs'), 'B Titr')
            cover_run.font.size = Pt(18)
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            
            
            document.add_paragraph()
            document.add_paragraph()
            set_run_language(cover_run)
            
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            footer_para = document.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_bidi_alignment(footer_para)
            footer_run = footer_para.add_run('شرکت دانش بنیان البرز افزار طبرستان (درنیکا)\nبهار 1404')
            footer_run.font.name = 'B Nazanin'
            footer_run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
            footer_run.font.size = Pt(12)
            set_run_language(footer_run)

            # برو به صفحه دوم
            document.add_page_break()

            cover_para = document.add_paragraph()
            cover_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_bidi_alignment(cover_para)
            cover_run = cover_para.add_run('فهرست مطالب')
            cover_run.font.name = 'B Titr'
            cover_run._element.rPr.rFonts.set(qn('w:cs'), 'B Titr')
            cover_run.font.size = Pt(14)

            add_table_of_contents(document)

            
            document.add_page_break()

            paragraph = document.add_paragraph()
            set_bidi_alignment(paragraph)
            text = f'مستند جداول دیتابیس: {db_name}'
            run = paragraph.add_run(text)
            run.font.name = "B Titr"
            run._element.rPr.rFonts.set(qn('w:cs'), 'B Titr')
            run.font.size = Pt(14)
            set_run_language(run)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            document.add_paragraph()
            if "doc_connection" not in st.session_state:
                        st.session_state.doc_connection = get_doc_db_connection()
            
            progress_bar = st.progress(0)
            progress_text = st.empty()
            table_count = len(st.session_state.tables)
            tbl_idx = 0
            for idx, table in enumerate(st.session_state.tables, start=1):
                
                table_name = table[0] if isinstance(table, (list, tuple)) else table
                #print(table_name)
                para = document.add_paragraph()
                set_bidi_alignment(para)

                query = """
                SELECT object_fa_name, object_description
                FROM table_dictionary
                WHERE object_name = %s
                AND type = 1
                AND (
                        is_general = 1
                        OR (
                            is_general = 0
                            AND table_name = %s
                            AND db_name = %s
                            AND db_type = %s
                        )
                    )
                ORDER BY is_general DESC
                LIMIT 1;
                """

                params = (table_name, table_name, db_name, db_type)
                with st.session_state.doc_connection.cursor() as doc_cursor:
                    doc_cursor.execute(query, params)
                    result = doc_cursor.fetchone()

                if result:
                    tpn = result[0]
                    tpd = result[1]
                   
                else:
                    tpn = ""
                    tpd = ""
                #document = Document()
                paragraph = document.add_paragraph(style = "Heading 1")
                set_bidi_alignment(paragraph)
                ttext = f" {idx} جدول {tpn} ({table_name}):"
                trun = paragraph.add_run(ttext)
                trun.font.name = "B Nazanin"
                trun._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                trun.font.size = Pt(14)
                set_run_language(trun)

                paragraph = document.add_paragraph()
                set_bidi_alignment(paragraph)
                dtext = f" {tpd}"
                drun = paragraph.add_run(dtext)
                drun.font.name = "B Nazanin"
                drun._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                drun.font.size = Pt(14)
                set_run_language(drun)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                table_docx = document.add_table(rows=1, cols=5)
                table_docx.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_docx.autofit = False
                total_width = 16  # عرض کل جدول به سانتی‌متر

                # درصدها به ترتیب ستون‌ها: [30%, 15%, 15%, 20%, 20%]
                column_widths = [0.30, 0.20, 0.10, 0.25, 0.15]
                for i, width_percent in enumerate(column_widths):
                    table_docx.cell(0, i).width = Cm(total_width * width_percent)
                table_docx.style = 'Table Grid'
                hdr_cells = table_docx.rows[0].cells
                headers = ["توضیحات", "ویژگی", "نوع", "فیلد فارسی نام", "فیلد نام"]
                
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in hdr_cells[i].paragraphs[0].runs:
                        run.font.name = 'B Titr'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Titr')
                        run.font.size = Pt(12)
                
                
                columns = get_table_columns(cursor, table_name)
                for col in columns:
                    col_name, data_type, is_nullable, default, REFERENCED_TABLE_NAME = col

                    row = table_docx.add_row().cells
                    cell = row[4]
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    text1 = col_name
                    for word in text1.split(" "):
                        run = para.add_run(word + " ")
                        if is_persian_or_digit(word):
                            run.font.name = "B Nazanin"
                            run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                            run.font.size = Pt(12)
                        else:
                            run.font.name = "Arial"
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                            run.font.size = Pt(10)
                    #set_cell_font(row[3], "")

                    query = """
                    SELECT object_fa_name, object_description
                    FROM table_dictionary
                    WHERE object_name = %s
                    AND type = 2
                    AND (
                            is_general = 1
                            OR (
                                is_general = 0
                                AND table_name = %s
                                AND db_name = %s
                                AND db_type = %s
                            )
                        )
                    ORDER BY is_general DESC
                    LIMIT 1;
                    """

                    params = (col_name, table_name, db_name, db_type)
                    #print(params)
                    with st.session_state.doc_connection.cursor() as doc_cursor:
                        doc_cursor.execute(query, params)
                        result = doc_cursor.fetchone()

                    if result:
                        fpn = result[0]
                        fpd = result[1]
                    else:
                        fpn = ""
                        fpd = ""

                    cell = row[3]
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    p = para._element
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    p.get_or_add_pPr().append(bidi)

                    run = para.add_run(fpn)
                    run.font.name = "B Nazanin"
                    run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(12)

                    cell = row[2]
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    text1 = data_type
                    for word in text1.split(" "):
                        run = para.add_run(word + " ")
                        if is_persian_or_digit(word):
                            run.font.name = "B Nazanin"
                            run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                            run.font.size = Pt(11)
                        else:
                            run.font.name = "Arial"
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                            run.font.size = Pt(9)
                    cell2 = row[1]
                    features = []
                    if is_nullable != "NO":
                        features.append("Nullable")
                    if default:
                        if features:
                            features.append("\n")
                            features.append(f"Default: {default}")
                        else:
                            features.append(f"Default: {default}")
                    if REFERENCED_TABLE_NAME:
                        if features:
                            features.append("\n")
                            features.append(f"References: {REFERENCED_TABLE_NAME}")
                        else:
                            features.append(f"References: {REFERENCED_TABLE_NAME}")
                        
                    cell2.text = "   ".join(features)
                    if features:

                        para = cell2.paragraphs[0]
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        text1 = cell2.text
                        para.clear()
                        for word in text1.split(" "):
                            run = para.add_run(word + " ")
                            if is_persian_or_digit(word):
                                run.font.name = "B Nazanin"
                                run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                                run.font.size = Pt(11)
                            else:
                                run.font.name = "Arial"
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                                run.font.size = Pt(9)
                    #set_cell_font(row[0], f"Default: {default}" if default else "")


                    cell = row[0]
                    para = cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    p = para._element
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    p.get_or_add_pPr().append(bidi)

                    run = para.add_run(fpd)
                    run.font.name = "B Nazanin"
                    run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(10)

                heading_para = document.add_paragraph()
                text = f'ارتباطات'
                set_bidi_alignment(heading_para)
                for word in text.split(" "):
                    run = heading_para.add_run(word + " ")
                    if is_persian_or_digit(word):
                        run.font.name = "B Nazanin"
                        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                        run.font.size = Pt(14)
                    else:
                        run.font.name = "Arial"
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                        run.font.size = Pt(12)
                heading_para.paragraph_format.space_after = 0 

                columns = get_foreign_key_join_clauses(cursor, table_name)
                if columns:
                    table_docx = document.add_table(rows=1, cols=3)
                    table_docx.alignment = WD_TABLE_ALIGNMENT.CENTER
                    #table_docx.autofit = False
                    total_width_cm = 15  # عرض کل جدول به سانتی‌متر

                    column_widths = [0.7, 0.15, 0.15]
                    for i, width_percent in enumerate(column_widths):
                        table_docx.cell(0, i).width = Cm(total_width_cm * width_percent)
                    table_docx.style = 'Table Grid'
                    hdr_cells = table_docx.rows[0].cells
                    headers = ["پیوند", "خارجی جدول", "اصلی جدول"]
                    
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in hdr_cells[i].paragraphs[0].runs:
                            run.font.name = 'B Titr'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Titr')
                            run.font.size = Pt(12)
                    
                    for col in columns:
                        foreign_table, primary_table, join_clause = col

                        row = table_docx.add_row().cells

                        # ستونی: اصلی جدول
                        para1 = row[2].paragraphs[0]
                        para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for word in primary_table.split(" "):
                            run = para1.add_run(word + " ")
                            if is_persian_or_digit(word):
                                run.font.name = "B Nazanin"
                                run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                                run.font.size = Pt(12)
                            else:
                                run.font.name = "Arial"
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                                run.font.size = Pt(10)

                        # ستونی: خارجی جدول
                        para2 = row[1].paragraphs[0]
                        para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for word in foreign_table.split(" "):
                            run = para2.add_run(word + " ")
                            if is_persian_or_digit(word):
                                run.font.name = "B Nazanin"
                                run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                                run.font.size = Pt(11)
                            else:
                                run.font.name = "Arial"
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                                run.font.size = Pt(9)

                        # ستونی: پیوند
                        para3 = row[0].paragraphs[0]
                        para3.clear()  # پاک‌سازی محتوای قبلی
                        para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for word in join_clause.split(" "):
                            run = para3.add_run(word + " ")
                            if is_persian_or_digit(word):
                                run.font.name = "B Nazanin"
                                run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                                run.font.size = Pt(11)
                            else:
                                run.font.name = "Arial"
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                                run.font.size = Pt(9)
                else:
                    heading_para = document.add_paragraph()
                    text = f'»ارتباطی یافت نشد«'
                    #set_bidi_alignment(heading_para)

                    p2 = heading_para._element
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    p2.get_or_add_pPr().append(bidi)

                    run = heading_para.add_run(text)

                    run.font.name = "B Nazanin"
                    run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(14)
                        
            

                heading_para = document.add_paragraph()
                text = f'کلیدهای منحصربفرد'
                set_bidi_alignment(heading_para)
                for word in text.split(" "):
                    run = heading_para.add_run(word + " ")
                    if is_persian_or_digit(word):
                        run.font.name = "B Nazanin"
                        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                        run.font.size = Pt(14)
                    else:
                        run.font.name = "Arial"
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                        run.font.size = Pt(12)
                heading_para.paragraph_format.space_after = 0 
                table_docx = document.add_table(rows=2, cols=2)
                table_docx.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_docx.autofit = False
                total_width = 16  # عرض کل جدول به سانتی‌متر

                column_widths = [0.5, 0.5]
                for i, width_percent in enumerate(column_widths):
                    table_docx.cell(0, i).width = Cm(total_width * width_percent)
                table_docx.style = 'Table Grid'
                hdr_cells = table_docx.rows[0].cells
                headers = ["توضیحات", "ستون نام"]
                
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in hdr_cells[i].paragraphs[0].runs:
                        run.font.name = 'B Titr'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Titr')
                        run.font.size = Pt(12)
                
                values = ["اصلی کلید", "id"]
                row_cells = table_docx.rows[1].cells
                for i, value in enumerate(values):
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for word in value.split(" "):
                        run = para.add_run(word + " ")
                        if is_persian_or_digit(word):
                            run.font.name = "B Nazanin"
                            run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                            run.font.size = Pt(12)
                        else:
                            run.font.name = "Arial"
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                            run.font.size = Pt(10)
                document.add_paragraph()
                tbl_idx = tbl_idx + 1
                percent = int((tbl_idx/table_count)*100)
                
                progress_text.write(f"مستند جدول {table_name} ساخته شد. ({tbl_idx}/{table_count}) - {percent}%")
                progress_bar.progress(percent)
            file_name = f"{db_name}.docx"
            document.save(file_name)



            with open(file_name, "rb") as file:
                st.download_button(
                    label="📥 دانلود فایل Word",
                    data=file,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.write("مستند دیتابیس ساخته شد!")

        if st.button("خروج"):
            st.session_state.clear() 
            st.rerun() 

        if 'show_selectbox' not in st.session_state:
            st.session_state.show_selectbox = False

        if st.button("ورود اطلاعات"):
            st.session_state.show_selectbox = True

        if st.session_state.show_selectbox:
            if ensure_main_connection():
                conn = st.session_state.connection
                cursor = conn.cursor()

                if "tables" not in st.session_state:
                    cursor.execute("SHOW TABLES")
                    st.session_state.tables = cursor.fetchall()

                table_names = [t[0] for t in st.session_state.tables]
                selected_table = st.selectbox("یک جدول انتخاب کنید:", sorted(table_names), key="table_select")
                st.session_state.selected_table = selected_table
                db_name = st.session_state.db_name
                db_type = st.session_state.db_type
                try:
                    if "doc_connection" not in st.session_state:
                        st.session_state.doc_connection = get_doc_db_connection()

                    doc_cursor2 = st.session_state.doc_connection.cursor()

                    
                    # چک کردن وجود رکورد عمومی
                    fetch_query = """
                        (
                            SELECT object_fa_name, object_description, is_general 
                            FROM table_dictionary
                            WHERE object_name = %s AND type = 1 AND is_general = 1
                            LIMIT 1
                        )
                        UNION
                        (
                            SELECT object_fa_name, object_description, is_general 
                            FROM table_dictionary
                            WHERE object_name = %s AND table_name = %s AND db_name = %s AND db_type = %s AND type = 1 AND is_general = 0
                            LIMIT 1
                        )
                        LIMIT 1
                    """
                    params = (st.session_state.selected_table, st.session_state.selected_table, st.session_state.selected_table, db_name, db_type)
                    fetch_values = (st.session_state.selected_table,)
                    doc_cursor2.execute(fetch_query, params)
                    record = doc_cursor2.fetchone()
                    

                    if record:
                        fa_name, description, is_generalt = record[0], record[1], record[2]
                    else:
                        fa_name, description, is_generalt = "", "", 0

                    
                    doc_cursor2.close()

                except Exception as e:
                    st.error(f"خطا در واکشی اطلاعات: {e}")
                    fa_name, description, is_generalt = "", "", 0
                
                if st.session_state.get("last_selected_table") != st.session_state.selected_table:
                    st.session_state.fa_name = fa_name
                    st.session_state.table_desc = description
                    st.session_state.table_public = bool(is_generalt)
                    st.session_state.last_selected_table = st.session_state.selected_table                

                if st.session_state.selected_table:
                    st.markdown("### اطلاعات جدول")
                    st.text_input("نام جدول:", value=st.session_state.selected_table, disabled=True)
                    table_fa_name = st.text_input("نام فارسی جدول:",  key="fa_name")
                    table_description = st.text_area("توضیح جدول:",  key="table_desc")
                    table_is_public = st.checkbox("آیا جدول عمومی است؟",  key="table_public")

                    if st.button(f"ثبت {st.session_state.selected_table}", key=f"submit_{st.session_state.selected_table}"):
                        is_general = 1 if table_is_public else 0
                        object_name = st.session_state.selected_table
                        object_fa_name = table_fa_name
                        object_description = table_description
                        table_name = object_name
                        db_name = st.session_state.db_name
                        db_type = st.session_state.db_type
                        type_value = 1
                        updated = False
                        try:
                            if "doc_connection" not in st.session_state:
                                st.session_state.doc_connection = get_doc_db_connection()

                            doc_cursor = st.session_state.doc_connection.cursor()
                                

                            check_query = """
                                (
                                SELECT object_fa_name, object_description, is_general 
                                FROM table_dictionary
                                WHERE object_name = %s AND type = 1 AND is_general = 1
                                LIMIT 1
                                )
                                UNION
                                (
                                    SELECT object_fa_name, object_description, is_general 
                                    FROM table_dictionary
                                    WHERE object_name = %s AND table_name = %s AND db_name = %s AND db_type = %s AND type = 1 AND is_general = 0
                                    LIMIT 1
                                )
                                LIMIT 1
                            """
                            params = (st.session_state.selected_table, st.session_state.selected_table, st.session_state.selected_table, db_name, db_type)
                            doc_cursor.execute(check_query, params)
                            existing = doc_cursor.fetchone()

                            if existing:
                                record_id = existing[0]
                                # بروزرسانی رکورد موجود
                                update_query = """
                                    UPDATE table_dictionary
                                    SET object_fa_name = %s,
                                        object_description = %s,
                                        table_name = %s,
                                        db_name = %s,
                                        db_type = %s
                                    WHERE object_name = %s
                                """
                                update_values = (
                                    object_fa_name, object_description,
                                    table_name, db_name, db_type,
                                    object_name
                                )

                                doc_cursor.execute(update_query, update_values)
                                st.session_state.doc_connection.commit()
                                if doc_cursor.rowcount > 0:
                                    updated = True
                                    #st.success("رکورد عمومی قبلی با موفقیت بروزرسانی شد ✅")
                                else:
                                    st.warning("رکورد عمومی یافت شد اما تغییری ایجاد نشد.")
                                    updated = True
       
                            if not updated:
                                insert_query = """
                                    INSERT INTO table_dictionary 
                                    (object_name, object_fa_name, object_description, table_name, db_name, db_type, is_general, type) 
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                                """
                                values = (
                                    object_name, object_fa_name, object_description,
                                    table_name, db_name, db_type, is_general, type_value
                                )
                                doc_cursor.execute(insert_query, values)
                                st.session_state.doc_connection.commit()
                                doc_cursor.close()
                            if updated:
                                st.success("رکورد عمومی قبلی با موفقیت بروزرسانی شد ✅")
                            else:
                                st.success("اطلاعات با موفقیت ثبت شد ✅")
                        except Exception as e:
                            st.error(f"خطا در ثبت اطلاعات جدول: {e}")

                    st.markdown("---")






                    st.markdown("### اطلاعات فیلدها")

                    columns = get_table_columns(cursor, st.session_state.selected_table)
                    for col in columns:
                        col_name = col[0]
                        data_type = col[1]
                        type_value = 2
                        st.markdown(f"**{col_name}: {data_type}**")

                        try:
                            if "doc_connection" not in st.session_state:
                                st.session_state.doc_connection = get_doc_db_connection()

                            doc_cursor2 = st.session_state.doc_connection.cursor()

                            
                            # چک کردن وجود رکورد عمومی
                            fetch_query = """
                                                                (
                                SELECT object_fa_name, object_description, is_general 
                                FROM table_dictionary
                                WHERE object_name = %s AND type = 2 AND is_general = 1
                                LIMIT 1
                                )
                                UNION
                                (
                                    SELECT object_fa_name, object_description, is_general 
                                    FROM table_dictionary
                                    WHERE object_name = %s AND table_name = %s AND db_name = %s AND db_type = %s AND type = 2 AND is_general = 0
                                    LIMIT 1
                                )
                                LIMIT 1
                            """
                            fetch_values = (col_name, type_value)
                            params = (col_name, col_name, st.session_state.selected_table, db_name, db_type)
                            #print(params)
                            doc_cursor2.execute(fetch_query, params)
                            record = doc_cursor2.fetchone()
                            

                            if record:
                                fa_name, description, is_generalf = record[0], record[1], record[2]
                            else:
                                fa_name, description, is_generalf = "", "", 0

                           

                            doc_cursor2.close()

                        except Exception as e:
                            st.error(f"خطا در واکشی اطلاعات: {e}")
                            fa_name, description, is_generalf = "", "", 0

                        

                        fa_name = st.text_input(f"نام فارسی فیلد ({col_name}):", value=fa_name, key=f"{selected_table}_{col_name}_fa")
                        description = st.text_area(f"توضیح فیلد ({col_name}):", value=description, key=f"{selected_table}_{col_name}_desc")
                        is_public = st.checkbox("آیا فیلد عمومی است؟", value=is_generalf, key=f"{selected_table}_{col_name}_public")

                        if st.button(f"ثبت {col_name}", key=f"submit_{col_name}"):
                            #st.write(f"اطلاعات برای فیلد {col_name} ثبت شد!")
                            
                            is_generalf = 1 if is_public else 0





                            updated = False
                            try:
                                if "doc_connection" not in st.session_state:
                                    st.session_state.doc_connection = get_doc_db_connection()

                                doc_cursor = st.session_state.doc_connection.cursor()
                                
                                check_query = """
                                (
                                SELECT object_fa_name, object_description, is_general 
                                FROM table_dictionary
                                WHERE object_name = %s AND type = 2 AND is_general = 1
                                LIMIT 1
                                )
                                UNION
                                (
                                    SELECT object_fa_name, object_description, is_general 
                                    FROM table_dictionary
                                    WHERE object_name = %s AND table_name = %s AND db_name = %s AND 
                                    db_type = %s AND type = 2 AND is_general = 0
                                    LIMIT 1
                                )
                                LIMIT 1
                                """
                                doc_cursor.execute(check_query, params)
                                params = (col_name, col_name, st.session_state.selected_table, db_name, db_type)
                                #print(params)
                                #doc_cursor2.execute(fetch_query, params)
                                existing = doc_cursor.fetchone()

                                if existing:
                                    record_id = existing[0]
                                    # بروزرسانی رکورد موجود
                                    update_query = """
                                        UPDATE table_dictionary
                                        SET object_fa_name = %s,
                                            object_description = %s   
                                        WHERE  table_name = %s and
                                            db_name = %s and
                                            db_type = %s and object_name = %s
                                    """
                                    update_values = (
                                        fa_name, description,
                                        st.session_state.selected_table, db_name, db_type,
                                        col_name
                                    )
                                    doc_cursor.execute(update_query, update_values)
                                    if doc_cursor.rowcount > 0:
                                        updated = True
                                    else:
                                        st.warning("رکورد عمومی یافت شد اما تغییری ایجاد نشد.")
                                        updated = True
                                        
                                        
                                        
                                if not updated:
                                    insert_query = """
                                        INSERT INTO table_dictionary 
                                        (object_name, object_fa_name, object_description, table_name, db_name, db_type, is_general, type) 
                                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                                    """
                                    values = (
                                        col_name, fa_name, description,
                                        st.session_state.selected_table, db_name, db_type, is_generalf, type_value
                                    )
                                    #print(values)
                                    doc_cursor.execute(insert_query, values)
                                st.session_state.doc_connection.commit()
                                doc_cursor.close()
                                if updated:
                                    st.success("رکورد عمومی قبلی با موفقیت بروزرسانی شد ✅")
                                else:
                                    st.success("اطلاعات با موفقیت ثبت شد ✅")
                            except Exception as e:
                                st.error(f"خطا در ثبت اطلاعات جدول: {e}")


                        st.markdown("---")


    else:
        st.error("اتصال به دیتابیس برقرار نیست!")

# شروع برنامه
if not st.session_state.get("db_connected"):
    db_connection_form()
else:
    db_tables_page()
